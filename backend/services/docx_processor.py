import logging
import json
import os
import asyncio
import random
import re
from pathlib import Path
from typing import Dict, Optional, List
from docx import Document
from io import BytesIO
from sqlalchemy.orm import Session
from google import genai
from google.genai import types
from backend.services.entity_utils import run_async

logger = logging.getLogger(__name__)


# Siglas técnicas que sí pueden aparecer con pocas palabras
SIGLAS_PERMITIDAS = {
    'sql', 'pl/sql', 'rest', 'soap', 'api', 'apis', 'erp', 'crm', 'scm',
    'rad', 'ddd', 'itil', 'bpmn', 'uml', 'scrum', 'devops', 'iot',
    'matlab', 'python', 'java', 'c#', 'c++', 'php', 'js', 'json', 'xml',
    'html', 'css', 'saas', 'paas', 'iaas', 'vm', 'vdi', 'bi', 'ai',
    'ml', 'dl', 'nlp', 'rpa', 'soa', 'nosql', 'mongodb', 'oracle',
    'mysql', 'postgresql', 'sqlite', 'aws', 'azure', 'gcp', 'docker',
    'kubernetes', 'git', 'cobit', 'iso', 'pmbok', 'cobra', 'cmmi',
    'http', 'https', 'tcp/ip', 'ip', 'lan', 'wan', 'vpn', 'wifi',
    'bluetooth', 'nfc', 'rfid', 'ldap', 'kerberos', 'oauth', 'jwt',
    'sdlc', 'oop', 'poo', 'mvc', 'mvvm', 'mvp', 'solid', 'dry', 'kiss',
    'yagni', 'tdd', 'bdd', 'crc', 'xp', 'kanban', 'lean',
    'rup', 'togaf', 'zachman', 'archimate', 'bsc', 'kpi', 'okr',
    'roi', 'van', 'tir', 'olap', 'oltp', 'etl', 'elt', 'hadoop',
    'spark', 'kafka', 'rabbitmq', 'mqtt', 'coap', 'ddl', 'dml',
    'dcl', 'tcl', 'acid', 'cap', 'sns', 'sqs', 'bd', 'sgbd', 'dbms'
}

# Palabras que por sí solas son demasiado genéricas para ser un tema útil
PALABRAS_GENERICAS = {
    'introducción', 'conceptos', 'conceptos generales', 'definición', 'definiciones',
    'fundamentos', 'componentes', 'elementos', 'aspectos', 'características',
    'tipos', 'clasificación', 'ventajas', 'desventajas', 'aplicaciones',
    'ejemplos', 'ejercicios', 'práctica', 'prácticas', 'taller', 'laboratorio',
    'casos', 'caso de estudio', 'problemas', 'problemas de aplicación',
    'generalidades', 'objetivos', 'resultados', 'conclusiones', 'resumen',
    'revisión', 'análisis', 'diseño', 'implementación', 'desarrollo',
    'administración', 'gestión', 'control', 'planificación', 'organización',
    'dirección', 'supervisión', 'evaluación', 'seguimiento', 'monitoreo',
    'conceptos básicos', 'aspectos básicos', 'noción', 'noción básica',
    'definición de', 'concepto de', 'elementos de', 'tipos de', 'clasificación de',
    'características de', 'ventajas de', 'desventajas de', 'aplicaciones de',
    'ejemplos de', 'ejercicios de', 'casos de', 'problemas de'
}

# Items administrativos, evaluativos o de socialización
PALABRAS_EXCLUIR = re.compile(
    r'(?i)('
    r'evaluaci[oó]n\s*(parcial|final|de\s*proceso|sustitutoria?)?|'
    r'examen\s*(parcial|final|sustitutorio)?|'
    r'retroalimentaci[oó]n\s*(de\s*contenidos?|de\s*los\s*contenidos?|y\s*nivelaci[oó]n|de\s*contenidos\s*desarrollados?)?|'
    r'nivelaci[oó]n|recuperaci[oó]n\s+(acad[eé]mica|de\s*contenidos)|'
    r'hito\s*\d(\s*del\s*proyecto)?|'
    r'semana\s*de\s*actualizaci[oó]n|foro\s*participativo|'
    r'presentaci[oó]n\s*final\s*del\s*proyecto|'
    r'informe\s*final|preparaci[oó]n\s*para|'
    r'presentaci[oó]n\s*del\s*(curso|s[ií]labo|programa)|'
    r'socializaci[oó]n\s*del\s*s[ií]labo|bienvenida(\s*al\s*curso)?|'
    r'clase\s*conferencia|datos\s*(personales|profesionales)\s*del\s*docente|'
    r'revisi[oó]n\s*del\s*reglamento|expectativas\s+laborales|'
    r'formaci[oó]n\s*de\s*equipos?|conformaci[oó]n\s*de\s*equipos?|'
    r'sustentaci[oó]n(\s*individual|\s*de\s*equipos?|\s*de\s*cada\s*integrante)|'
    r'presentaci[oó]n(\s*y\s*revisi[oó]n)?\s*de\s*trabajos|'
    r'continuaci[oó]n\s*de\s*sustentaciones|diagn[oó]stico\s*de\s*recursos|'
    r'accion\s*tutorial|tutor[ií]a\s+acad[eé]mica|'
    r'resultados\s*de\s*la\s*investigaci[oó]n\s*y\s*su\s*impacto|'
    r'desarrollo\s+de\s+una\s+clase\s+conferencia|'
    r'presentaci[oó]n\s+formal\s+de\s+trabajos|'
    r'contenidos\s+trabajados\s+hasta|la\s+semana\s+\d+|los\s+casos\s+pertinentes|'
    r'observaci[oó]n\s+e\s+interrogantes|respuestas\s+a\s+preguntas|'
    r'avance\s+de\s+la\s+asignatura|avance\s+de\s+trabajo|'
    r'\(\s*[A-Z]{2,4}\s*\)\s*\d+%'
    r')'
)

# Actividades del estudiante/docente, no contenido temático
ACTIVIDADES_EXCLUIR = re.compile(
    r'(?i)('
    r'^(se\s+(desarrolla|desarrollan|revisa|revisan|presenta|presentan|conforman|'
    r'gu[ií]a|da|sustenta|discute|identifican|caracteriza|caracterizan|'
    r'elabora|elaboran|aplica|aplican|define|definen|realiza|realizan|'
    r'internaliza|internalizan|complementa|complementan|revisan|desarrolla|desarrollan))\b|'
    r'^(trabajo\s+en\s+grupo\s+para|caso\s+de\s+aplicaci[oó]n|'
    r'en\s+la\s+pr[aá]ctica|en\s+clase\s+te[oó]rica|'
    r'desarrollo\s+de\s+contenidos|se\s+revisan\s+las\s+propuestas|'
    r'aplicaci[oó]n\s+de\s+test|desarrollo\s+del\s+test|'
    r'continuaci[oó]n\s+de\s+sustentaciones|'
    r'presentaci[oó]n\s+y\s+revisi[oó]n\s+de\s+trabajos|'
    r'sustentaci[oó]n\s+de\s+cada\s+integrante|'
    r'aplicaci[oó]n\s+de\s+test\s+y\s+retroalimentaci[oó]n|'
    r'revisi[oó]n\s+de\s+devops|revisi[oó]n\s+de\s+las\s+propuestas|'
    r'el\s+desarrollo\s+de\s+sistemas\s+software|'
    r'respuestas\s+a\s+preguntas|'
    r'entregado\s+en\s+plataforma\s+canvas|'
    r'trabajos\s+individuales\s+no\s+son\s+v[aá]lidos)|'
    r'objeto\s+de\s+estudio|'
    r'organizaci[oó]n\s+objeto\s+de\s+estudio|'
    r'monitoreado\s+por\s+el\s+docente|'
    r'en\s+el\s+avance\s+de\s+la\s+asignatura|'
    r'actividades\s+o\s+acciones\s+de\s+retroalimentaci[oó]n'
    r')'
)

# Conectores al inicio que indican item huérfano
CONECTORES_HUERFANOS = re.compile(
    r'^(según|segun|de|y|e|o|u|en|para|por|con|sin|sobre|bajo|entre|durante|'
    r'mediante|tales\s+como|como|además|también|por\s+lo\s+tanto|es\s+decir|'
    r'esto\s+es|o\s+sea|principalmente|especialmente|particularmente|'
    r'incluyendo|incluye|incluido|tales|etc|etc\.)\b',
    re.IGNORECASE
)


def prompt_temas_silabo(texto_silabo: str) -> str:
    return f"""
Eres un extractor especializado de contenidos temáticos de sílabos universitarios.
Extraes ÚNICAMENTE lo que está en la columna "Contenidos Temáticos" de la tabla de
programación semanal. No inventas, no condensas y no agregas información externa.

TAREA:
Revisa semana por semana la columna "Contenidos Temáticos" y devuelve una lista
de temas. Cada tema debe ser una frase con sentido completo.

REGLAS OBLIGATORIAS:
1. Cada tema debe tener al menos 3 palabras con significado técnico. No aceptes
   palabras sueltas como "Introducción", "Definición", "Componentes", "Tipos",
   "Aplicaciones", "Reportes", "Ventajas".
2. NO dividas una frase por cada punto. Divide SOLO cuando haya bullets explícitos
   (•, -, *, números) o saltos de línea claros dentro de una misma celda.
3. Si una celda dice "Apis. Servicios Rest. Servicios SOAP", devuelve un tema
   completo como "Diseño e integración de APIs REST y SOAP", no items sueltos.
4. NO incluyas presentaciones del curso, socializaciones del sílabo, retroalimentaciones,
   evaluaciones, hitos de proyecto, foros, semanas de actualización académica ni
   presentación de proyectos.
5. NO incluyas actividades del estudiante como "Se desarrolla un caso...",
   "Trabajo en grupo para...", "Caso de aplicación...".
6. Si una semana no tiene contenido temático real, omítela por completo.
7. No dupliques temas idénticos o casi idénticos.

EJEMPLOS DE ENTRADA Y SALIDA:

Entrada:
Semana 1 | Introducción • Conceptos generales • Objetivos del administrador de la base de datos

Salida:
[
  "Objetivos del administrador de la base de datos"
]

Entrada:
Semana 10 | Apis. Servicios Rest. Servicios SOAP

Salida:
[
  "Diseño e integración de APIs REST y SOAP"
]

Entrada:
Semana 4 | RETROALIMENTACIÓN DE LOS CONTENIDOS TRABAJADOS

Salida:
[]

FORMATO DE SALIDA:
Devuelve ÚNICAMENTE este JSON:
{{
  "nombre": "string",
  "codigo": "string",
  "ciclo": 0,
  "temas": ["string"]
}}

Si no hay temas extraíbles: "temas": [].

DOCUMENTO:
{texto_silabo}
"""


def _es_sigla_permitida(texto: str) -> bool:
    """Permite siglas técnicas conocidas aunque tengan pocas palabras."""
    limpio = texto.lower().strip().rstrip('.:,;')
    if limpio in SIGLAS_PERMITIDAS:
        return True
    palabras = limpio.split()
    if palabras and palabras[-1].rstrip('.:,;') in SIGLAS_PERMITIDAS:
        return True
    return False


def _es_generico(texto: str) -> bool:
    """Detecta items que son puro relleno genérico."""
    limpio = texto.lower().strip().rstrip('.:,;')
    if limpio in PALABRAS_GENERICAS:
        return True
    palabras = limpio.split()
    if len(palabras) <= 2 and limpio in PALABRAS_GENERICAS:
        return True
    patrones_relleno = [
        r'^(introducci[oó]n|conceptos?|definiciones?|fundamentos|componentes|'
        r'elementos|aspectos|caracter[ií]sticas|tipos|clasificaci[oó]n|'
        r'ventajas|desventajas|aplicaciones|ejemplos|ejercicios|pr[aá]cticas?|'
        r'problemas|generalidades|objetivos|revisi[oó]n|an[aá]lisis|diseño|'
        r'implementaci[oó]n|desarrollo|administraci[oó]n|gesti[oó]n|control|'
        r'planificaci[oó]n|organizaci[oó]n|direcci[oó]n|supervisi[oó]n|'
        r'seguimiento|monitoreo)\s*(al|a|de|del|en|sobre|para|por)?\s*$'
    ]
    for patron in patrones_relleno:
        if re.search(patron, limpio, re.IGNORECASE):
            return True
    return False


def _normalizar_texto(texto: str) -> str:
    """Limpia espacios, saltos de línea y puntuación redundante."""
    t = texto.replace('\n', ' ').replace('\r', ' ').replace('\t', ' ')
    t = re.sub(r'\s+', ' ', t).strip()
    t = re.sub(r'\s+([.,;:!?])', r'\1', t)
    t = re.sub(r'([.,;:!?])\s*,\s*', r'\1 ', t)
    return t


def _limpiar_prefijos(texto: str) -> str:
    """Quita bullets, puntos y guiones residuales al inicio de un item."""
    t = texto.strip()
    # \u00B7 = middle dot, común en sílabos
    t = re.sub(r'^[\s\.•\-–—*+\u00B7]+', '', t).strip()
    return t


def _dividir_celda_en_temas(texto_celda: str) -> List[str]:
    """
    Divide una celda de contenidos temáticos en items individuales.
    Respeta bullets, saltos de línea y oraciones separadas por punto+mayúscula.
    """
    if not texto_celda:
        return []

    t = texto_celda
    # Reemplazar bullets unicode por •
    t = re.sub(r'[\u2022\u2023\u25E6\u2043\u2219\u25AA\u25AB\u2013\u2014]', '•', t)
    # Reemplazar guiones/numeración al inicio de línea por •
    t = re.sub(r'(^|\n)\s*[-*+]\s+', r'\1• ', t)
    t = re.sub(r'(^|\n)\s*\d+[\.\)]\s+', r'\1• ', t)

    # Separar por bullets o saltos de línea
    partes = re.split(r'\s*•\s*|\n+', t)

    temas = []
    for parte in partes:
        parte = _normalizar_texto(parte)
        parte = _limpiar_prefijos(parte)
        if not parte:
            continue

        # Si la parte tiene múltiples oraciones, dividir por punto+espacio+mayúscula
        oraciones = re.split(r'\.\s+(?=[A-ZÁÉÍÓÚÑ])', parte)
        for oracion in oraciones:
            oracion = _normalizar_texto(oracion)
            oracion = _limpiar_prefijos(oracion)
            oracion = oracion.rstrip('.')
            if oracion:
                temas.append(oracion)

    return temas


def _fusionar_huerfanos(temas: List[str]) -> List[str]:
    """Fusiona items huérfanos (cortos que empiezan con conector) con el anterior."""
    if not temas:
        return []
    result = [temas[0]]
    for t in temas[1:]:
        palabras = t.split()
        if len(palabras) <= 3 and CONECTORES_HUERFANOS.match(t):
            result[-1] = result[-1] + ' ' + t[0].lower() + t[1:] if len(t) > 1 else result[-1] + ' ' + t
            result[-1] = _normalizar_texto(result[-1])
        else:
            result.append(t)
    return result


def _es_actividad(texto: str) -> bool:
    """Detecta si un item es una actividad del estudiante/docente, no contenido temático."""
    if ACTIVIDADES_EXCLUIR.search(texto):
        return True
    return False


def limpiar_temas(temas: List[str]) -> List[str]:
    """
    Limpieza robusta de temas extraídos:
    - Normaliza espacios y puntuación
    - Elimina vacíos y duplicados
    - Filtra items administrativos/evaluativos
    - Filtra actividades del estudiante
    - Filtra items muy cortos o genéricos
    - Fusiona items huérfanos
    """
    if not temas:
        return []

    vistos = set()
    filtrados = []

    for t in temas:
        t = _normalizar_texto(t)
        if not t:
            continue

        # Filtro administrativo/evaluativo
        if PALABRAS_EXCLUIR.search(t):
            logger.info(f"Tema excluido (admin/eval): '{t[:80]}...'")
            continue

        # Filtro de actividades
        if _es_actividad(t):
            logger.info(f"Tema excluido (actividad): '{t[:80]}...'")
            continue

        # Filtro de genéricos puros
        if _es_generico(t):
            logger.info(f"Tema excluido (genérico): '{t[:80]}...'")
            continue

        # Filtro de longitud: mínimo 3 palabras, salvo siglas
        palabras = t.split()
        if len(palabras) < 3 and not _es_sigla_permitida(t):
            logger.info(f"Tema excluido (corto): '{t[:80]}...'")
            continue

        # Filtro de items huérfanos por conector
        if len(palabras) <= 3 and CONECTORES_HUERFANOS.match(t) and not _es_sigla_permitida(t):
            logger.info(f"Tema excluido (huérfano): '{t[:80]}...'")
            continue

        # Evitar duplicados exactos (case-insensitive, sin puntuación final)
        clave = re.sub(r'[^a-z0-9áéíóúñ]+', '', t.lower())
        if clave in vistos:
            logger.info(f"Tema duplicado descartado: '{t[:80]}...'")
            continue
        vistos.add(clave)

        filtrados.append(t)

    # Fusionar items huérfanos con el anterior
    filtrados = _fusionar_huerfanos(filtrados)

    return filtrados


def _detectar_columnas_tabla(fila_encabezado: List[str]) -> tuple:
    """Detecta los índices de semana y contenido temático en una fila de encabezado."""
    idx_semana = None
    idx_contenido = None
    for i, cell_text in enumerate(fila_encabezado):
        txt_lower = cell_text.lower()
        if any(p in txt_lower for p in ['semana', 'n° semanas', 'n semanas', 'semanas']):
            idx_semana = i
        if any(p in txt_lower for p in ['contenidos temáticos', 'contenido temático', 'contenidos', 'contenido tematico']):
            idx_contenido = i
    return idx_semana, idx_contenido


def _extraer_temas_manualmente(docx_bytes: bytes) -> List[str]:
    """
    Extrae contenidos temáticos de la tabla de programación semanal de forma determinista.
    No requiere llamadas a IA.
    """
    try:
        doc = Document(BytesIO(docx_bytes))
    except Exception as e:
        logger.error(f"Error abriendo DOCX para extracción manual: {e}")
        return []

    temas_crudos = []
    idx_semana_global = None
    idx_contenido_global = None
    ultima_semana = None
    semanas_procesadas = set()

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)

        if not rows:
            continue

        # Detectar encabezado en la primera fila
        idx_semana, idx_contenido = _detectar_columnas_tabla(rows[0])
        if idx_contenido is not None:
            idx_semana_global = idx_semana if idx_semana is not None else 0
            idx_contenido_global = idx_contenido
            data_rows = rows[1:]
        elif idx_contenido_global is not None:
            idx_semana = idx_semana_global
            idx_contenido = idx_contenido_global
            data_rows = rows
        else:
            if len(rows[0]) >= 2:
                idx_semana = 0
                idx_contenido = 1
                data_rows = rows
            else:
                continue

        for cells in data_rows:
            if idx_contenido >= len(cells):
                continue

            texto_semana = cells[idx_semana].strip() if idx_semana is not None and idx_semana < len(cells) else ""
            texto_contenido = cells[idx_contenido].strip()

            # Detectar si esta fila inicia una nueva semana
            semana_match = re.search(r'semana\s*(\d+)', texto_semana, re.IGNORECASE)
            es_continuacion = False
            if semana_match:
                num_semana = int(semana_match.group(1))
                ultima_semana = num_semana
                es_continuacion = False
            elif ultima_semana is not None:
                es_continuacion = True
            else:
                continue

            if not texto_contenido:
                continue

            # Para filas de continuación, solo procesar si no hemos visto este contenido
            # para la semana actual (evita duplicados de actividades)
            if es_continuacion:
                clave_semana_contenido = f"{ultima_semana}_{texto_contenido.lower()}"
                if clave_semana_contenido in semanas_procesadas:
                    continue
                semanas_procesadas.add(clave_semana_contenido)
            else:
                clave_semana_contenido = f"{ultima_semana}_{texto_contenido.lower()}"
                if clave_semana_contenido in semanas_procesadas:
                    continue
                semanas_procesadas.add(clave_semana_contenido)

            items = _dividir_celda_en_temas(texto_contenido)
            temas_crudos.extend(items)

    return limpiar_temas(temas_crudos)


class DOCXProcessor:
    def __init__(self):
        self.project_id = os.getenv("GOOGLE_CLOUD_PROJECT")
        self.location = os.getenv("GOOGLE_CLOUD_LOCATION", "global")
        
        try:
            self.client = genai.Client(vertexai=True, project=self.project_id, location=self.location)
            logger.info(f"genai SDK inicializado en el proyecto {self.project_id} para DOCX")
        except Exception as e:
            logger.error(f"Error iniciando genai Client en DOCX: {e}", exc_info=True)
            self.client = None

    async def llamar_flash_lite(self, texto: str) -> dict:
        if not self.client: return {}
        import time
        max_retries = 3
        for attempt in range(max_retries):
            try:
                if attempt > 0:
                    wait = 5 * (2 ** (attempt - 1)) + random.uniform(0, 2)
                    logger.warning(f"⏳ [Sílabo] Flash Lite retry {attempt+1}/{max_retries} en {wait:.1f}s...")
                    await asyncio.sleep(wait)
                start_time = time.time()
                response = await asyncio.wait_for(
                    self.client.aio.models.generate_content(
                        model="gemini-3.1-flash-lite",
                        contents=prompt_temas_silabo(texto),
                        config=types.GenerateContentConfig(
                            response_mime_type="application/json",
                            temperature=0.1
                        )
                    ),
                    timeout=120.0
                )
                elapsed = time.time() - start_time
                logger.info(f"⏱️ [Sílabo] Flash Lite tardó: {elapsed:.2f} segundos")
                try:
                    return json.loads(response.text.strip())
                except Exception as e:
                    logger.error(f"Error parseando Flash Lite (DOCX): {e}")
                    return {}
            except Exception as e:
                if ("429" in str(e) or "RESOURCE_EXHAUSTED" in str(e)) and attempt < max_retries - 1:
                    continue
                logger.error(f"Error en Flash Lite Sílabo (intento {attempt+1}): {e}", exc_info=True)
                return {}

    async def procesar_silabo(self, texto_silabo: str) -> dict:
        return await self.llamar_flash_lite(texto_silabo)

    def extract_text_from_docx(self, docx_bytes: bytes) -> str:
        """Extrae todo el texto plano del DOCX, incluyendo tablas, respetando el orden real."""
        try:
            from docx.oxml.ns import qn
            doc = Document(BytesIO(docx_bytes))
            text = []

            for block in doc.element.body:
                # Párrafo normal
                if block.tag == qn('w:p'):
                    from docx.text.paragraph import Paragraph
                    p = Paragraph(block, doc)
                    if p.text.strip():
                        text.append(p.text)

                # Tabla
                elif block.tag == qn('w:tbl'):
                    from docx.table import Table
                    table = Table(block, doc)
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text.append(" | ".join(row_text))

            return "\n".join(text)
        except Exception as e:
            logger.error(f"Error leyendo DOCX crudo: {e}", exc_info=True)
            return ""

    def _extraer_metadata_desde_filename(self, filename: str) -> Dict:
        """Extrae nombre y código aproximados desde el nombre del archivo."""
        base = filename.replace('.docx', '').replace('_', ' ').strip()
        match = re.search(r'\b([A-Z]{2,6}-\d{3,4})\b', base)
        codigo = match.group(1) if match else None
        nombre = base
        if codigo:
            partes = base.split(codigo)
            if len(partes) > 1:
                nombre = partes[-1].strip(' -')
        return {'nombre': nombre, 'codigo': codigo}

    def extract_syllabus_info(self, docx_bytes: bytes, filename: str = "") -> Dict:
        try:
            # 1. Leer texto crudo
            full_text = self.extract_text_from_docx(docx_bytes)
            if not full_text:
                return {'success': False, 'error': 'DOCX vacío o ilegible'}

            logger.info(f"Extrayendo temas semanales del sílabo: {filename}")

            # 2. Extracción determinista desde la tabla semanal
            temas_manuales = _extraer_temas_manualmente(docx_bytes)
            logger.info(f"Temas extraídos manualmente: {len(temas_manuales)}")

            # 3. Fallback con Flash Lite si la extracción manual es muy pobre
            ai_data = None
            temas_gemini = []
            if len(temas_manuales) < 5:
                logger.info("Extracción manual pobre, intentando con Gemini Flash Lite...")
                ai_data = run_async(self.procesar_silabo(full_text))
                if ai_data:
                    temas_gemini = limpiar_temas(ai_data.get('temas', []))
                    logger.info(f"Temas extraídos por Gemini: {len(temas_gemini)}")

            # 4. Combinar y limpiar
            todos_temas = list(temas_manuales)
            if temas_gemini:
                vistos = {re.sub(r'[^a-z0-9áéíóúñ]+', '', t.lower()) for t in todos_temas}
                for t in temas_gemini:
                    clave = re.sub(r'[^a-z0-9áéíóúñ]+', '', t.lower())
                    if clave not in vistos:
                        todos_temas.append(t)
                        vistos.add(clave)

            temas_limpios = limpiar_temas(todos_temas)
            logger.info(f"Total temas finales: {len(temas_limpios)}")

            # 5. Metadata
            metadata_filename = self._extraer_metadata_desde_filename(filename) if filename else {}
            nombre = None
            codigo = None
            ciclo = 1

            if ai_data:
                nombre = ai_data.get('nombre')
                codigo = ai_data.get('codigo')
                ciclo = ai_data.get('ciclo', 1)

            if not nombre:
                nombre = metadata_filename.get('nombre')
            if not nombre and filename:
                nombre = filename.replace('.docx', '').replace('_', ' ')
            if not codigo:
                codigo = metadata_filename.get('codigo')

            return {
                'success': True,
                'nombre': nombre,
                'codigo': codigo,
                'ciclo': ciclo,
                'temas': temas_limpios,
                'raw_text': full_text,
                'raw_text_length': len(full_text)
            }
        except Exception as e:
            logger.error(f"Error procesando sílabo: {e}", exc_info=True)
            return {'success': False, 'error': str(e)}

    def save_curso_to_db(self, db: Session, syllabus_info: Dict, drive_file_id: str) -> Optional[int]:
        try:
            from backend.database import crud
            existing = crud.get_curso_by_drive_id(db, drive_file_id)

            data = {
                "nombre": syllabus_info.get('nombre', 'Curso Desconocido'),
                "codigo": syllabus_info.get('codigo'),
                "ciclo": int(syllabus_info.get('ciclo', 1)),
                "temas": syllabus_info.get('temas', []),
                # Deprecated: se mantienen vacíos por compatibilidad
                "entidades_clave": [],
                "competencias_tecnicas": ""
            }

            if existing:
                crud.update_curso(db, existing.id, **data)
                return existing.id
            else:
                curso = crud.create_curso(db, drive_file_id=drive_file_id, **data)
                return curso.id
        except Exception as e:
            logger.error(f"Error BD Curso: {e}", exc_info=True)
            try:
                db.rollback()
            except Exception as rollback_err:
                logger.error(f"Error en rollback: {rollback_err}")
            return None

docx_processor = DOCXProcessor()
